"""
Professional Documentation Generator and Exporter
Analyzes codebase and generates standardized documentation with watermarks
"""

import os
import re
import ast
import logging
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image, ListFlowable, ListItem, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import markdown
import time

class DocumentationGenerator:
    """Handles document generation in various formats from markdown content"""
    
    def __init__(self):
        """Initialize with default styles and configurations"""
        self.styles = getSampleStyleSheet()
        
        # Define all custom styles we want to use
        custom_styles = {
            'Code': ParagraphStyle(
                name='Code',
                parent=self.styles['Normal'],
                fontName='Courier',
                fontSize=9,
                backColor=colors.lightgrey,
                spaceAfter=12
            ),
            'CodeBlock': ParagraphStyle(
                name='CodeBlock',
                parent=self.styles['Normal'],
                fontName='Courier',
                fontSize=9,
                backColor=colors.lightgrey,
                leftIndent=12,
                rightIndent=12,
                spaceAfter=12,
                spaceBefore=12
            )
        }
        
        # Safely register or update each style
        for style_name, style in custom_styles.items():
            if style_name in self.styles:
                # Update existing style with new properties
                existing_style = self.styles[style_name]
                for attr, value in style.__dict__.items():
                    if attr != 'name' and hasattr(existing_style, attr):
                        setattr(existing_style, attr, value)
            else:
                # Add new style
                self.styles.add(style)
        
    def export_to_pdf(self, content: str, output_path: str) -> str:
        """
        Convert markdown content to a professionally formatted PDF with watermark
        
        Args:
            content: Markdown content to convert
            output_path: Where to save the PDF file
            
        Returns:
            str: Path to the generated PDF file
        """
        try:
            # Function for page templates with minimal, professional footers
            def add_page_elements(canvas, doc):
                # Add subtle footer with page number only (except on first page)
                canvas.saveState()
                if canvas.getPageNumber() > 1:  # Skip footer on first/cover page
                    canvas.setFont('Helvetica', 8)
                    canvas.setFillColor(colors.grey)
                    page_text = f"Page {canvas.getPageNumber()}"
                    canvas.drawRightString(doc.width + 50, 20, page_text)
                canvas.restoreState()
                
                # No header for cleaner look, more content space
            
            # Set up document with optimized margins for more content area
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=54,  # Reduced margins
                leftMargin=54,
                topMargin=56,  # Reduced top margin
                bottomMargin=56  # Reduced bottom margin
            )
            

            
            # Create document story
            story = []
            
            # Extract title for cover page (first H1 heading)
            title = "Code Documentation"
            title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
            if title_match:
                title = title_match.group(1)
            
            # Add professional title with optimized spacing
            story.append(Spacer(1, 0.5*inch))  # Balanced top spacing
            
            # Professional title with improved styling
            story.append(Paragraph(title, ParagraphStyle(
                name="Title", 
                parent=self.styles['Title'],
                fontSize=26,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#2c3e50"),  # Professional dark blue
                leading=32,  # Control line height
                spaceAfter=8
            )))
            
            # Thin separator line for professional look
            separator = Table([[""]], colWidths=[3*inch], rowHeights=[1])
            separator.setStyle(TableStyle([
                ('LINEABOVE', (0, 0), (-1, -1), 0.5, colors.HexColor("#95a5a6")),
            ]))
            story.append(separator)
            
            # Generation date in subtle gray
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 
                         ParagraphStyle(
                             name="Subtitle", 
                             parent=self.styles['Normal'],
                             alignment=TA_CENTER,
                             fontSize=10,
                             textColor=colors.HexColor("#7f8c8d"),  # Subtle gray
                             spaceAfter=4
                         )))
            
            # Add professionally styled table of contents
            story.append(Spacer(1, 0.3*inch))  # Better spacing before TOC
            
            # TOC header with professional styling
            toc_entry = Paragraph("Table of Contents", ParagraphStyle(
                name="TOC_Header",
                parent=self.styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor("#34495e"),
                spaceBefore=0,
                spaceAfter=8,
                borderPadding=0
            ))
            story.append(toc_entry)
            
            # Subtle line below TOC header
            toc_separator = Table([[""]], colWidths=[4*inch], rowHeights=[1])
            toc_separator.setStyle(TableStyle([
                ('LINEABOVE', (0, 0), (-1, -1), 0.3, colors.HexColor("#bdc3c7")),
            ]))
            story.append(toc_separator)
            story.append(Spacer(1, 0.1*inch))
            
            # Extract TOC entries (only if there are headings)
            toc_entries = []
            for match in re.finditer(r'^(#+)\s+(.+)$', content, re.MULTILINE):
                level = len(match.group(1))
                heading_text = match.group(2)
                if level <= 3:  # Only include up to H3 in TOC
                    indent = '    ' * (level - 1)
                    toc_entries.append(f"{indent}• {heading_text}")
            
            # Add professionally styled TOC entries
            if toc_entries:
                # Create TOC table for consistent spacing and better appearance
                toc_data = []
                for entry in toc_entries:
                    # Get indentation level from leading spaces
                    indent_level = 0
                    for i, char in enumerate(entry):
                        if char != ' ':
                            indent_level = i // 4
                            break
                    
                    # Remove bullet point and trim
                    entry_text = entry.strip().lstrip('• ')
                    
                    # Add styled TOC entry with consistent formatting
                    toc_data.append([Paragraph(
                        entry, 
                        ParagraphStyle(
                            name=f"TOC_Entry_{indent_level}",
                            parent=self.styles['Normal'],
                            fontSize=10,
                            leftIndent=15*indent_level,
                            spaceBefore=3,  # Tight but readable spacing
                            spaceAfter=3,
                            textColor=colors.HexColor("#4e5b60" if indent_level == 0 else "#5d6d7e")
                        )
                    )])
                
                # Create TOC table with optimized layout
                toc_table = Table(toc_data, colWidths=[doc.width - 100])
                toc_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                    ('TOPPADDING', (0, 0), (-1, -1), 1),
                ]))
                story.append(toc_table)
                
                # Only add page break if TOC is large enough to justify separate page
                # Otherwise keep content flowing on same page for shorter documents
                if len(content.strip()) > 0 and len(toc_entries) > 5:
                    story.append(Spacer(1, 0.3*inch))
                    story.append(PageBreak())
                else:
                    # For short TOCs, add visual separator
                    story.append(Spacer(1, 0.3*inch))
                    separator = Table([[""]], colWidths=[doc.width - 108], rowHeights=[1])
                    separator.setStyle(TableStyle([
                        ('LINEABOVE', (0, 0), (-1, -1), 0.3, colors.HexColor("#ecf0f1")),
                    ]))
                    story.append(separator)
                    story.append(Spacer(1, 0.2*inch))
            
            # Helper functions and content processing
            current_list_items = []
            
            def flush_list():
                """Helper to add accumulated list items to story with professional styling"""
                nonlocal current_list_items, story
                if current_list_items:
                    # Create a professional compact list style
                    list_style = ParagraphStyle(
                        name="ListItem",
                        parent=self.styles['Normal'],
                        fontSize=10,  # Slightly smaller for lists
                        leading=14,   # Tight but readable line spacing
                        spaceBefore=0,
                        spaceAfter=0,
                        leftIndent=10,
                    )
                    
                    # Group list items with same indentation level for better spacing
                    indent_groups = {}
                    for item in current_list_items:
                        # Count leading spaces to determine indentation level
                        indent_match = re.match(r'^(\s*)', item)
                        indent_level = len(indent_match.group(1)) if indent_match else 0
                        
                        if indent_level not in indent_groups:
                            indent_groups[indent_level] = []
                        indent_groups[indent_level].append(item.strip())
                    
                    # Process each indentation group
                    for indent_level, items in sorted(indent_groups.items()):
                        bullets = ListFlowable(
                            [ListItem(Paragraph(item, list_style)) for item in items],
                            bulletType='bullet',
                            start='•',
                            leftIndent=10 + (indent_level * 10),  # Progressive indentation
                            rightIndent=0,
                            bulletFontName='Helvetica',
                            bulletFontSize=9,
                            bulletOffsetY=0,
                            spaceBefore=2 if indent_level == 0 else 0,  # Space only before top-level lists
                            spaceAfter=4 if indent_level == 0 else 2,    # Less space after nested lists
                        )
                        story.append(bullets)
                    
                    current_list_items = []
            
            # Split content into lines for processing
            lines = content.split('\n')
            
            # Track list and code block state
            in_code = False
            code_lines = []
            code_language = ""
            
            for line in lines:
                if line.startswith('```'):
                    # Toggle code block state
                    if in_code:
                        # End code block
                        flush_list()  # Flush any pending list items
                        code_text = '\n'.join(code_lines)
                        
                        # Create a styled code block with border and background
                        if code_language:
                            # Add language label within the code block header
                            code_header = Table([[Paragraph(f"<i>Language: {code_language}</i>",
                                         ParagraphStyle(
                                             name="CodeLabel",
                                             parent=self.styles['Normal'],
                                             fontSize=8,
                                             textColor=colors.darkgrey
                                         ))]],
                                    colWidths=[doc.width - 72])
                            code_header.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                                ('TOPPADDING', (0, 0), (-1, -1), 2),
                                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                                ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.grey),
                            ]))
                            story.append(code_header)
                        
                        # Code content with more compact style and explicit width control
                        code_style = ParagraphStyle(
                            name="CodeBlockOptimized",
                            parent=self.styles['CodeBlock'],
                            fontSize=8.5,  # Slightly smaller font
                            leading=10,  # Tighter line spacing
                            spaceBefore=0,
                            spaceAfter=0
                        )
                        
                        # Calculate available width (accounting for page margins)
                        available_width = doc.width - 72
                        
                        # Create preformatted text with optimized style
                        code_preformatted = Preformatted(code_text, code_style)
                        
                        # Create a table with tight margins and borders
                        code_container = Table([[code_preformatted]],
                                            colWidths=[available_width])
                        code_container.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, -1), colors.whitesmoke),  # Lighter background
                            ('BOX', (0, 0), (-1, -1), 0.5, colors.lightgrey),  # Lighter border
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),  # Reduced padding
                            ('TOPPADDING', (0, 0), (-1, -1), 3),  # Reduced padding
                            ('LEFTPADDING', (0, 0), (-1, -1), 8),  # Reduced padding
                            ('RIGHTPADDING', (0, 0), (-1, -1), 8),  # Reduced padding
                        ]))
                        story.append(code_container)
                        story.append(Spacer(1, 4))  # Reduced spacing
                        
                        code_lines = []
                        code_language = ""
                    else:
                        # Start code block - check for language
                        code_language = line[3:].strip()
                    in_code = not in_code
                    continue
                    
                if in_code:
                    code_lines.append(line)
                    continue
                
                # Handle lists
                list_match = re.match(r'^(\s*)[*-]\s+(.+)$', line)
                if list_match:
                    text = list_match.group(2)
                    # Process inline code in list items
                    text = re.sub(r'`([^`]+)`', 
                              lambda m: f'<font face="Courier" size="9" color="#222222">{m.group(1)}</font>', 
                              text)
                    current_list_items.append(text)
                    continue
                else:
                    flush_list()  # End list if we're not on a list item
                
                # Handle headings
                if line.startswith('#'):
                    level = len(re.match(r'^#+', line).group())
                    text = line.lstrip('#').strip()
                    
                    # Add spacing before headings (except first heading)
                    if story and not isinstance(story[-1], PageBreak):
                        story.append(Spacer(1, 0.2 * inch))
                    
                    # Use modern, professional heading styles with careful spacing
                    if level == 1:
                        style = ParagraphStyle(
                            name="CustomH1",
                            parent=self.styles['Heading1'],
                            fontSize=18,
                            textColor=colors.HexColor("#2c3e50"),  # Dark blue-gray
                            fontName="Helvetica-Bold",
                            spaceBefore=15,
                            spaceAfter=10,
                            borderPadding=0,
                            leading=22  # Controlled line height
                        )
                        # Add subtle underline for H1 sections
                        if story and not isinstance(story[-1], PageBreak):
                            h1_separator = Table([[""]], colWidths=[doc.width - 108], rowHeights=[1])
                            h1_separator.setStyle(TableStyle([
                                ('LINEBELOW', (0, 0), (-1, -1), 0.3, colors.HexColor("#e0e6ed")),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                            ]))
                            story.append(h1_separator)
                    elif level == 2:
                        style = ParagraphStyle(
                            name="CustomH2",
                            parent=self.styles['Heading2'],
                            fontSize=15,
                            textColor=colors.HexColor("#34495e"),  # Medium blue-gray
                            fontName="Helvetica-Bold",
                            spaceBefore=12,
                            spaceAfter=8,
                            leading=18  # Controlled line height
                        )
                    else:
                        style = ParagraphStyle(
                            name=f"CustomH{level}",
                            parent=self.styles[f'Heading{min(level, 4)}'],
                            fontSize=13 - (level-3),  # Smaller for deeper headings
                            textColor=colors.HexColor("#52656b"),  # Subtle blue-gray
                            spaceBefore=10,
                            spaceAfter=6,
                            leading=16  # Controlled line height
                        )
                    
                    story.append(Paragraph(text, style))
                    continue
                    
                # Handle regular paragraphs
                if line.strip():
                    # Process inline code
                    line = re.sub(r'`([^`]+)`', 
                              lambda m: f'<font face="Courier" size="9" color="#222222" backcolor="#f5f5f5">{m.group(1)}</font>', 
                              line)
                    
                    # Process bold text
                    line = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', line)
                    
                    # Process italic text
                    line = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', line)
                    
                    story.append(Paragraph(line, self.styles['Normal']))
                elif story and not isinstance(story[-1], Spacer) and not isinstance(story[-1], PageBreak):
                    # Add space between paragraphs (avoid duplicate spacers)
                    story.append(Spacer(1, 6))
            
            # Flush any remaining list items
            flush_list()
            
            # Build and save the PDF with page templates
            doc.build(story, onFirstPage=add_page_elements, onLaterPages=add_page_elements)
            return output_path
            
        except ImportError as e:
            raise ImportError(f"Failed to create PDF: Required package not installed. Install reportlab package. Error: {str(e)}")
        except Exception as e:
            raise Exception(f"Failed to create PDF: {str(e)}")
            
    def export_to_docx(self, content: str, output_path: str) -> str:
        """
        Convert markdown content to a professionally formatted DOCX with styling and watermark
        
        Args:
            content: Markdown content to convert  
            output_path: Where to save the DOCX file
            
        Returns:
            str: Path to the generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor, Mm
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
            from docx.enum.dml import MSO_THEME_COLOR_INDEX
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Create document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = "Code Documentation"
            doc.core_properties.author = "Code-Docgen Enterprise"
            doc.core_properties.created = datetime.now()
            
            # Configure document styles for professional appearance
            styles = doc.styles
            
            # Set professional page margins
            for section in doc.sections:
                section.top_margin = Mm(13)       # Optimized professional top margin
                section.bottom_margin = Mm(13)    # Optimized professional bottom margin
                section.left_margin = Mm(19)      # Optimized professional left margin
                section.right_margin = Mm(19)     # Optimized professional right margin
            
            # Create professional typography styles for better readability
            # Normal text style
            normal_style = styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(10)
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(2)
            normal_style.paragraph_format.line_spacing = 1.08  # Slightly tighter line spacing
            
            # Heading styles for professional hierarchy
            for i in range(1, 4):  # Headings 1-3
                if f'Heading {i}' in styles:
                    heading_style = styles[f'Heading {i}']
                    heading_style.font.name = 'Calibri'
                    heading_style.font.size = Pt(16 - (i*2))  # 16pt for H1, 14pt for H2, 12pt for H3
                    heading_style.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)  # Professional blue-gray
                    heading_style.paragraph_format.space_before = Pt(12 - (i*2))
                    heading_style.paragraph_format.space_after = Pt(6)
            
            # Add title
            title = "Code Documentation"
            title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
            if title_match:
                title = title_match.group(1)
                
            # Create professional cover page with consistent branding
            # Substantial top margin for professional title placement
            doc.add_paragraph().add_run().add_break()
            
            # Professional title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(24)
            title_run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)  # Professional dark blue
            
            # Add horizontal line for professional appearance
            border_para = doc.add_paragraph()
            border_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            border_para.paragraph_format.space_before = Pt(12)
            border_para.paragraph_format.space_after = Pt(12)
            border_run = border_para.add_run()
            border_run.add_text('_' * 40)  # 40-character line
            border_run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)  # Light gray
            
            # Add date in professional format
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.paragraph_format.space_before = Pt(6)
            date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            date_run.italic = True
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)  # Subtle gray
            
            # Watermark removed
            
            # Add section break
            doc.add_page_break()
            
            # Add professionally styled table of contents
            toc_header = doc.add_heading("Table of Contents", level=1)
            toc_header.runs[0].font.color.rgb = RGBColor(0x34, 0x49, 0x5e)  # Professional blue-gray
            
            # Add subtle separator line below TOC header
            toc_sep = doc.add_paragraph()
            toc_sep.paragraph_format.space_before = Pt(0)
            toc_sep.paragraph_format.space_after = Pt(6)
            toc_sep_run = toc_sep.add_run("_" * 40)  # Simple underline using text
            toc_sep_run.font.color.rgb = RGBColor(0xbd, 0xc3, 0xc7)  # Light gray
            
            # Extract and add TOC entries with optimized spacing and professional styling
            toc_entries = []
            toc_count = 0
            for match in re.finditer(r'^(#+)\s+(.+)$', content, re.MULTILINE):
                level = len(match.group(1))
                heading_text = match.group(2)
                if level <= 3:  # Only include up to H3 in TOC
                    # Create TOC entry with professional styling
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.2 * (level - 1))  # Slightly tighter indentation
                    para.paragraph_format.space_before = Pt(1 if level > 1 else 2)
                    para.paragraph_format.space_after = Pt(1 if level > 1 else 2)
                    
                    # Use varying bullet styles by level for professional hierarchy
                    if level == 1:
                        bullet = "•"  # Filled circle for level 1
                    elif level == 2:
                        bullet = "◦"  # Open circle for level 2
                    else:
                        bullet = "▪"  # Small square for level 3
                    
                    # Add entry with styled bullet and text
                    entry_run = para.add_run(f"{bullet} {heading_text}")
                    entry_run.font.size = Pt(11 - (level-1))  # Slightly smaller font for deeper levels
                    
                    if level == 1:
                        entry_run.font.bold = True
                        entry_run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)  # Dark blue for level 1
                    else:
                        entry_run.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)  # Medium blue for lower levels
                    
                    toc_count += 1
            
            # Add visual separator after TOC
            toc_end_sep = doc.add_paragraph()
            toc_end_sep.paragraph_format.space_before = Pt(6)
            toc_end_sep.paragraph_format.space_after = Pt(6)
            # Simple horizontal line
            toc_end_sep.add_run("_" * 20).font.color.rgb = RGBColor(0xec, 0xf0, 0xf1)  # Very light gray
            
            # Only add page break after TOC if TOC is substantial
            if toc_count > 5:
                doc.add_page_break()
            else:
                # Add some spacing after a short TOC for better visual separation
                doc.add_paragraph().paragraph_format.space_after = Pt(12)
            
            # Split content into lines
            lines = content.split('\n')
            
            # Track list state
            in_list = False
            list_items = []
            
            # Track code block state  
            in_code = False
            code_lines = []
            code_language = ""
            
            def add_list_items():
                """Helper to add accumulated list items"""
                nonlocal list_items
                if list_items:
                    for indent, text in list_items:
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.left_indent = Pt(18 * (indent // 2))
                        p.paragraph_format.space_before = Pt(0)  # Minimized spacing
                        p.paragraph_format.space_after = Pt(0)   # Minimized spacing
                        
                        # Process any formatting in the list item text
                        parts = re.split(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)', text)
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                # Inline code - optimize font size
                                code_run = p.add_run(part[1:-1])
                                code_run.font.name = 'Courier New'
                                code_run.font.size = Pt(8.5)  # Slightly smaller for code
                                code_run.font.color.rgb = RGBColor(51, 51, 51)
                                
                            elif part.startswith('**') and part.endswith('**'):
                                # Bold text
                                bold_run = p.add_run(part[2:-2])
                                bold_run.bold = True
                                
                            elif part.startswith('*') and part.endswith('*'):
                                # Italic text
                                italic_run = p.add_run(part[1:-1])
                                italic_run.italic = True
                                
                            elif part:
                                # Regular text
                                p.add_run(part)
                    
                    list_items = []
            
            for line in lines:
                if line.startswith('```'):
                    if in_code:
                        # End code block
                        # Add language identifier if present
                        if code_language:
                            lang_para = doc.add_paragraph()
                            lang_para.paragraph_format.space_before = Pt(6)
                            lang_para.paragraph_format.space_after = Pt(0)
                            lang_run = lang_para.add_run(f"Language: {code_language}")
                            lang_run.italic = True
                            lang_run.font.size = Pt(8)
                            lang_run.font.color.rgb = RGBColor(100, 100, 100)
                        
                        # Create professionally styled code block
                        code_text = '\n'.join(code_lines)
                        
                        # Create a table for the code block with background shading
                        # This provides better control over spacing and appearance than paragraph shading
                        code_table = doc.add_table(rows=1, cols=1)
                        code_table.style = 'Table Grid'  # Light borders
                        
                        # Set cell properties for code block
                        code_cell = code_table.cell(0, 0)
                        code_cell.width = Inches(6.0)  # Set width to fit page
                        
                        # Add shading to cell
                        cell_properties = code_cell._element.tcPr
                        shading_obj = OxmlElement('w:shd')
                        shading_obj.set(qn('w:fill'), "F8F8F8")  # Light gray background
                        cell_properties.append(shading_obj)
                        
                        # Add code paragraph with optimized spacing
                        code_para = code_cell.paragraphs[0]  # First paragraph in cell
                        code_para.paragraph_format.space_before = Pt(3)
                        code_para.paragraph_format.space_after = Pt(3)
                        code_para.paragraph_format.line_spacing = 1.0  # Single spacing for code
                        
                        # Add code text with professional monospace styling
                        code_run = code_para.add_run(code_text)
                        code_run.font.name = 'Consolas'  # More readable monospace font
                        code_run.font.size = Pt(8.5)  # Slightly smaller for compactness
                        code_run.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
                        
                        # Add minimal spacing after the table
                        spacer = doc.add_paragraph()
                        spacer.paragraph_format.space_before = Pt(0)
                        spacer.paragraph_format.space_after = Pt(3)
                        
                        # Reset code block tracking
                        code_lines = []
                        code_language = ""
                    else:
                        # Start code block - check for language
                        code_language = line[3:].strip()
                    in_code = not in_code
                    continue
                
                if in_code:
                    code_lines.append(line)
                    continue
                
                # Handle headings
                if line.startswith('#'):
                    # Add any pending list items
                    add_list_items()
                    
                    level = len(re.match(r'^#+', line).group())
                    text = line.lstrip('#').strip()
                    doc.add_heading(text, level=min(level, 9))
                    continue
                
                # Handle lists
                list_match = re.match(r'^(\s*)[*-]\s+(.+)$', line)
                if list_match:
                    indent = len(list_match.group(1))
                    text = list_match.group(2)
                    list_items.append((indent, text))
                    in_list = True
                    continue
                elif in_list:
                    # End list
                    add_list_items()
                    in_list = False
                
                # Handle regular paragraphs
                if line.strip():
                    # Add any pending list items
                    add_list_items()
                    
                    p = doc.add_paragraph()
                    
                    # Process inline code
                    parts = re.split(r'(`[^`]+`)', line)
                    for part in parts:
                        if part.startswith('`') and part.endswith('`'):
                            # Inline code
                            run = p.add_run(part[1:-1])  # Remove backticks
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
                        else:
                            # Regular text
                            p.add_run(part)
            
            # Add any remaining list items
            add_list_items()
                            
            # Save the document
            doc.save(output_path)
            return output_path
            
        except ImportError as e:
            raise ImportError(f"Failed to create DOCX: Required package not installed. Install python-docx package. Error: {str(e)}")
        except Exception as e:
            raise Exception(f"Failed to create DOCX: {str(e)}")

    def export_to_html(self, content: str, output_path: str) -> str:
        """
        Convert markdown content to a professionally styled HTML document
        
        Args:
            content: Markdown content to convert
            output_path: Where to save the HTML file
            
        Returns:
            str: Path to the generated HTML file
        """
        try:
            # Convert markdown to HTML
            html_content = markdown.markdown(
                content, 
                extensions=[
                    'markdown.extensions.fenced_code',
                    'markdown.extensions.tables',
                    'markdown.extensions.codehilite'
                ]
            )
            
            # Get document title
            title = "Code Documentation"
            title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
            if title_match:
                title = title_match.group(1)
            
            # Add syntax highlighting CSS and modern styling
            css = """
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.5; /* Tighter line spacing */
                color: #333;
                max-width: 900px;
                margin: 0 auto;
                padding: 15px; /* Reduced padding */
            }
            header {
                text-align: center;
                padding: 10px; /* Reduced padding */
                margin-bottom: 20px; /* Reduced margin */
                border-bottom: 1px solid #eee;
                position: relative;
            }
            /* No watermark */
            h1 {
                color: #2c3e50;
                margin-top: 20px; /* Reduced margin */
                margin-bottom: 15px; /* Controlled spacing */
            }
            h2 {
                color: #3498db;
                margin-top: 18px; /* Reduced margin */
                margin-bottom: 12px; /* Controlled spacing */
                border-bottom: 1px solid #eee;
                padding-bottom: 4px; /* Reduced padding */
            }
            h3, h4 {
                color: #2c3e50;
                margin-top: 15px; /* Reduced margin */
                margin-bottom: 10px; /* Controlled spacing */
            }
            pre {
                background-color: #f8f8f8;
                border: 1px solid #ddd;
                border-radius: 3px;
                padding: 10px; /* Reduced padding */
                margin: 10px 0; /* Reduced margin */
                overflow: auto;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 13px; /* Smaller font */
                line-height: 1.3; /* Tighter line spacing */
            }
            code {
                background-color: #f0f0f0;
                padding: 1px 3px; /* Reduced padding */
                border-radius: 2px;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 0.9em;
            }
            pre code {
                background-color: transparent;
                padding: 0;
            }
            blockquote {
                border-left: 4px solid #ddd;
                padding-left: 16px;
                margin-left: 0;
                color: #777;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }
            table, th, td {
                border: 1px solid #ddd;
            }
            th, td {
                padding: 12px;
                text-align: left;
            }
            th {
                background-color: #f8f8f8;
            }
            tr:nth-child(even) {
                background-color: #f8f8f8;
            }
            .footer {
                text-align: center;
                margin-top: 40px;
                padding-top: 20px;
                border-top: 1px solid #eee;
                color: #777;
                font-size: 0.9em;
            }
            .code-header {
                background-color: #e7e7e7;
                border: 1px solid #ddd;
                border-bottom: none;
                border-radius: 4px 4px 0 0;
                padding: 5px 16px;
                font-size: 0.8em;
                color: #777;
                display: flex;
                justify-content: space-between;
                margin-top: 20px;
            }
            .code-block-container {
                margin: 0;
            }
            .code-block-container + pre {
                border-top-left-radius: 0;
                border-top-right-radius: 0;
                margin-top: 0;
            }
            .toc {
                background-color: #f8f8f8;
                border: 1px solid #ddd;
                border-radius: 4px;
                padding: 16px;
                margin: 20px 0;
            }
            .toc ul {
                margin: 0;
                padding-left: 20px;
            }
            .toc-title {
                font-weight: bold;
                margin-bottom: 10px;
            }
            """
            
            # Process code blocks to add language labels
            def process_code_blocks(html):
                # Add language labels to code blocks
                pattern = r'<pre><code class="language-([^"]+)">(.*?)</code></pre>'
                replacement = r'<div class="code-header"><span>Language: \1</span></div><pre class="code-block-container"><code class="language-\1">\2</code></pre>'
                return re.sub(pattern, replacement, html, flags=re.DOTALL)
            
            # Extract TOC
            toc_html = '<div class="toc"><div class="toc-title">Table of Contents</div><ul>\n'
            for match in re.finditer(r'^(#+)\s+(.+)$', content, re.MULTILINE):
                level = len(match.group(1))
                heading_text = match.group(2)
                if level <= 3:  # Only include up to H3 in TOC
                    # Create a CSS-friendly ID from the heading text
                    heading_id = re.sub(r'[^a-z0-9]', '-', heading_text.lower())
                    toc_html += f'<li style="margin-left: {(level-1)*15}px"><a href="#{heading_id}">{heading_text}</a></li>\n'
            toc_html += '</ul></div>'
            
            # Add IDs to headings for TOC links
            def add_heading_ids(html):
                for match in re.finditer(r'<h([1-3])>(.*?)</h\1>', html):
                    heading_level = match.group(1)
                    heading_text = match.group(2)
                    heading_id = re.sub(r'[^a-z0-9]', '-', heading_text.lower())
                    html = html.replace(
                        match.group(0),
                        f'<h{heading_level} id="{heading_id}">{heading_text}</h{heading_level}>'
                    )
                return html
            
            # Process code blocks and add heading IDs
            html_content = process_code_blocks(html_content)
            html_content = add_heading_ids(html_content)
            
            # Construct the complete HTML document
            full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
    {css}
    </style>
</head>
<body>
    <header>
        <h1>{title}</h1>
        <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    </header>
    
    {toc_html}
    
    <div class="content">
    {html_content}
    </div>
    
    <div class="footer">
        <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    </div>
</body>
</html>"""
            
            # Write HTML to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
                
            return output_path
            
        except ImportError as e:
            raise ImportError(f"Failed to create HTML: Required package not installed. Install markdown package. Error: {str(e)}")
        except Exception as e:
            raise Exception(f"Failed to create HTML: {str(e)}")
    
    def export_to_txt(self, content: str, output_path: str) -> str:
        """
        Convert markdown content to a well-formatted plain text document
        
        Args:
            content: Markdown content to convert
            output_path: Where to save the TXT file
            
        Returns:
            str: Path to the generated TXT file
        """
        try:
            # Get document title
            title = "CODE DOCUMENTATION"
            title_match = re.search(r'^# (.+)$', content, re.MULTILINE)
            if title_match:
                title = title_match.group(1).upper()
            
            # Initialize output lines
            output_lines = []
            
            # Add header with borders (more compact)
            border = "=" * 80
            output_lines.append(border)
            output_lines.append(title.center(80))
            output_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}".center(80))
            # Removed branding line
            output_lines.append(border)
            output_lines.append("")  # Single blank line
            
            # Process content by lines
            lines = content.split('\n')
            in_code = False
            code_block = []
            
            for line in lines:
                # Handle code blocks
                if line.startswith('```'):
                    if in_code:
                        # End code block with formatted border
                        output_lines.append("-" * 80)
                        for code_line in code_block:
                            output_lines.append("    " + code_line)
                        output_lines.append("-" * 80)
                        output_lines.append("")
                        code_block = []
                    else:
                        # Start code block - get language if specified
                        lang = line[3:].strip()
                        if lang:
                            output_lines.append(f"[Code: {lang}]")
                    in_code = not in_code
                    continue
                
                if in_code:
                    code_block.append(line)
                    continue
                
                # Handle headings with different formatting levels
                if line.startswith('# '):
                    output_lines.append("")
                    text = line[2:].strip()
                    output_lines.append(text.upper())
                    output_lines.append("=" * len(text))
                    output_lines.append("")
                elif line.startswith('## '):
                    output_lines.append("")
                    text = line[3:].strip()
                    output_lines.append(text)
                    output_lines.append("-" * len(text))
                    output_lines.append("")
                elif line.startswith('### '):
                    output_lines.append("")
                    output_lines.append(line[4:].strip())
                    output_lines.append("")
                elif line.startswith('#### '):
                    output_lines.append("")
                    output_lines.append(line[5:].strip())
                    output_lines.append("")
                
                # Handle list items
                elif re.match(r'^\s*[*-]\s', line):
                    indent = len(re.match(r'^\s*', line).group())
                    text = re.sub(r'^\s*[*-]\s+', '', line)
                    # Clean markdown formatting
                    text = re.sub(r'`([^`]+)`', r'\1', text)  # Remove code marks
                    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold
                    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Remove italic
                    output_lines.append(" " * indent + "- " + text)
                
                # Regular paragraphs
                elif line.strip():
                    # Clean markdown formatting
                    line = re.sub(r'`([^`]+)`', r'\1', line)  # Remove code marks
                    line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # Remove bold
                    line = re.sub(r'\*([^*]+)\*', r'\1', line)  # Remove italic
                    
                    # Wrap text for better readability (max 80 chars)
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line) + len(word) + 1 > 80:
                            output_lines.append(current_line)
                            current_line = word
                        else:
                            if current_line:
                                current_line += " " + word
                            else:
                                current_line = word
                    
                    if current_line:
                        output_lines.append(current_line)
                
                # Add a blank line for empty markdown lines to preserve paragraph structure
                elif not line.strip():
                    output_lines.append("")
            
            # Add footer (more compact)
            output_lines.append("")
            output_lines.append(border)
            output_lines.append("End of documentation".center(80))
            output_lines.append(border)
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(output_lines))
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Failed to create TXT: {str(e)}")
    
    @staticmethod
    def create_temporary_file(content: str, export_format: str, filename: str = None) -> str:
        """
        Create a temporary file with the given content in the specified format.
        
        Args:
            content: The markdown content to export
            export_format: The format to export to ('txt', 'html', 'md', 'docx', 'pdf')
            filename: Optional custom filename
            
        Returns:
            str: Path to the created temporary file
            
        Raises:
            ImportError: If required packages are not installed
            Exception: If document generation fails
        """
        try:
            logger = logging.getLogger(__name__)
            logger.info(f"DocumentationGenerator: Starting temp file creation. Format: {export_format}")
            logger.info(f"DocumentationGenerator: Content length: {len(content) if content else 0}")
            logger.info(f"DocumentationGenerator: Using filename: {filename}")

            # Validate content
            if not content:
                raise ValueError("Content cannot be empty")
            if not isinstance(content, str):
                raise ValueError(f"Content must be a string, got {type(content)}")
                
            # Validate format
            if not export_format:
                raise ValueError("Export format is required")
            if not isinstance(export_format, str):
                raise ValueError(f"Export format must be a string, got {type(export_format)}")
            if export_format not in ['pdf', 'html', 'md', 'txt', 'docx']:
                raise ValueError(f"Unsupported format: {export_format}")

            # Create temp directory if it doesn't exist
            temp_dir = os.path.join('media', 'temp')
            os.makedirs(temp_dir, exist_ok=True)
            
            # Generate temporary filename
            if not filename:
                timestamp = int(time.time())
                filename = f"documentation_export_{timestamp}"
            elif not isinstance(filename, str):
                raise ValueError(f"Filename must be a string, got {type(filename)}")
        
            temp_filename = f"{filename}.{export_format}"
            temp_path = os.path.join(temp_dir, temp_filename)
        
            try:
                # Create generator with better error handling
                try:
                    generator = DocumentationGenerator()
                except Exception as init_error:
                    logger.error(f"Error initializing DocumentationGenerator: {str(init_error)}")
                    raise Exception(f"Failed to initialize document generator: {str(init_error)}")
                
                if export_format == 'pdf':
                    try:
                        return generator.export_to_pdf(content, temp_path)
                    except Exception as e:
                        logger.error(f"DocumentationGenerator: Error in PDF export - {str(e)}")
                        raise Exception(f"Failed to create PDF file: {str(e)}")
                    
                elif export_format == 'docx':
                    try:
                        return generator.export_to_docx(content, temp_path)
                    except Exception as e:
                        logger.error(f"DocumentationGenerator: Error in DOCX export - {str(e)}")
                        raise Exception(f"Failed to create DOCX file: {str(e)}")
                    
                elif export_format == 'html':
                    try:
                        return generator.export_to_html(content, temp_path)
                    except Exception as e:
                        logger.error(f"DocumentationGenerator: Error in HTML export - {str(e)}")
                        raise Exception(f"Failed to create HTML file: {str(e)}")
                        
                elif export_format == 'md':
                    try:
                        # Markdown files don't need conversion, just save directly
                        formatted_md = content
                        with open(temp_path, 'w', encoding='utf-8') as f:
                            f.write(formatted_md)
                        return temp_path
                    except Exception as e:
                        logger.error(f"DocumentationGenerator: Error in Markdown export - {str(e)}")
                        raise Exception(f"Failed to create Markdown file: {str(e)}")
                        
                elif export_format == 'txt':
                    try:
                        return generator.export_to_txt(content, temp_path)
                    except Exception as e:
                        logger.error(f"DocumentationGenerator: Error in TXT export - {str(e)}")
                        raise Exception(f"Failed to create TXT file: {str(e)}")
                    # Write content directly
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                
                return temp_path
                
            except ImportError as e:
                raise ImportError(f"Failed to create {export_format.upper()} file: Required package not installed. Error: {str(e)}")
            except Exception as e:
                raise Exception(f"Failed to create {export_format.upper()} file: {str(e)}")
        except Exception as e:
            logger.error(f"DocumentationGenerator: Error in temp file creation - {str(e)}")
            raise